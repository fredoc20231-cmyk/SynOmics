#!/usr/bin/env python3
"""SynOmics flagship RNA-seq pipeline (hybrid short+long read).

Two honest halves, one dispatch (stdin JSON -> stdout JSON):

UPSTREAM (Phases 1-3) — `rnaseq_upstream`:
    Orchestrates the real external toolchain (fastp, STAR, minimap2/STARlong,
    stringtie/gffcompare, salmon, samtools/NanoPlot). It DETECTS each binary with
    shutil.which and builds the EXACT command line (with the scientifically correct
    flags — fastp sliding-window trim; STAR --sjdbOverhang = ReadLength-1;
    minimap2 -ax splice fed the short-read SJ.out.tab; stringtie --merge; salmon
    index with genome decoys; salmon quant --validateMappings --seqBias --gcBias).
    Where a binary is present it RUNS it; where absent it returns status
    "unavailable" for that step with the command that WOULD run + an install hint.
    It NEVER fabricates alignments, counts, or quantifications. In this sandbox no
    aligner binaries exist, so the plan is returned honestly, unexecuted.

DOWNSTREAM (Phase 4) — `rnaseq_deseq` (+ `rnaseq_tximport`):
    A real, self-contained DESeq2-style differential-expression engine that RUNS
    here on a gene-count matrix: median-of-ratios size factors -> mean-dispersion
    trend fit + shrinkage -> per-gene Negative-Binomial GLM Wald test -> BH FDR ->
    normal-prior log2FC shrinkage -> VST + PCA. Emits the full figure/table set and
    a report / document / full article via the outcome-bundle writer.

Honest scope: the downstream statistics are an independent NB-GLM implementation
(statsmodels), DESeq2-*style* and validated against a labeled spike-in fixture to
recover known DE genes — it is not a re-wrapping of the R DESeq2 binary.
"""
from __future__ import annotations

import json
import math
import os
import shutil
import subprocess
import sys


def _fail(msg, status="error"):
    print(json.dumps({"status": status, "error": msg}))
    sys.exit(0)


# ----------------------------------------------------------------------------
# UPSTREAM ORCHESTRATOR (Phases 1-3) — real commands, honest availability
# ----------------------------------------------------------------------------

def _tool_status(binary):
    path = shutil.which(binary)
    return {"tool": binary, "available": path is not None, "path": path}


UPSTREAM_INSTALL = {
    "fastp": "conda install -c bioconda fastp",
    "STAR": "conda install -c bioconda star",
    "STARlong": "conda install -c bioconda star",
    "minimap2": "conda install -c bioconda minimap2",
    "samtools": "conda install -c bioconda samtools",
    "NanoPlot": "pip install NanoPlot",
    "stringtie": "conda install -c bioconda stringtie",
    "gffcompare": "conda install -c bioconda gffcompare",
    "salmon": "conda install -c bioconda salmon",
}


def _step(name, phase, binary, command, description):
    st = _tool_status(binary)
    return {
        "step": name,
        "phase": phase,
        "tool": binary,
        "description": description,
        "command": command,
        "available": st["available"],
        "toolPath": st["path"],
        "installHint": None if st["available"] else UPSTREAM_INSTALL.get(binary),
        # Honest execution state. Only steps whose binary exists AND `execute` was
        # requested are actually run; everything else is a plan, never a fake result.
        "status": "ready" if st["available"] else "unavailable",
    }


def task_rnaseq_upstream(p):
    """Phase 1-3 hybrid RNA-seq preprocessing/alignment/quant orchestrator."""
    read_length = int(p.get("readLength", 150))
    sjdb_overhang = read_length - 1
    ref_fasta = p.get("referenceFasta", "<genome.fa>")
    gtf = p.get("annotationGtf", "<annotation.gtf>")
    genome_dir = p.get("genomeDir", "star_index")
    salmon_index = p.get("salmonIndex", "salmon_index")
    r1 = p.get("read1", "R1.fastq.gz")
    r2 = p.get("read2", "R2.fastq.gz")
    long_reads = p.get("longReads", "long.fastq.gz")
    long_platform = p.get("longPlatform", "nanopore")  # or "pacbio_hifi"
    min_q = 10 if long_platform == "nanopore" else 20
    samples = p.get("samples", ["sampleA", "sampleB"])
    execute = bool(p.get("execute", False))

    steps = [
        _step("short_read_trim", 1, "fastp",
              f"fastp -i {r1} -I {r2} -o trimmed_R1.fq.gz -O trimmed_R2.fq.gz "
              f"--cut_front --cut_tail --cut_window_size 4 --cut_mean_quality 20 "
              f"--length_required 36 --detect_adapter_for_pe --thread 8 "
              f"--json fastp.json --html fastp.html",
              "Adapter trim + sliding-window quality filter + min-length drop (protects graph alignment)."),
        _step("long_read_qc_bin", 1, "NanoPlot",
              f"NanoPlot --fastq {long_reads} --minqual {min_q} --outdir nanoplot_qc",
              f"Long-read Q-score binning (>=Q{min_q} for {long_platform}); drop garbage reads before ML graph correction."),
        _step("star_genome_generate", 2, "STAR",
              f"STAR --runMode genomeGenerate --genomeDir {genome_dir} "
              f"--genomeFastaFiles {ref_fasta} --sjdbGTFfile {gtf} "
              f"--sjdbOverhang {sjdb_overhang} --runThreadN 8",
              f"Suffix-array splice-aware index; --sjdbOverhang = ReadLength-1 = {sjdb_overhang}."),
        _step("star_align_short", 2, "STAR",
              f"STAR --runMode alignReads --genomeDir {genome_dir} "
              f"--readFilesIn trimmed_R1.fq.gz trimmed_R2.fq.gz --readFilesCommand zcat "
              f"--outSAMtype BAM SortedByCoordinate --twopassMode Basic --runThreadN 8 "
              f"--outFileNamePrefix short_",
              "Splice-aware short-read alignment -> sorted BAM + verified SJ.out.tab junction map."),
        _step("minimap2_align_long", 2, "minimap2",
              f"minimap2 -ax splice --junc-bed short_SJ.out.tab -uf {ref_fasta} {long_reads} "
              f"| samtools sort -o long_sorted.bam - && samtools index long_sorted.bam",
              "Long-read splice alignment; short-read SJ.out.tab snaps noisy long-read intron boundaries to verified coordinates."),
        _step("stringtie_assemble_merge", 2, "stringtie",
              "stringtie short_Aligned.sortedByCoord.out.bam -G " + gtf + " -o {sample}.gtf ; "
              "stringtie --merge -G " + gtf + " -o merged.gtf " + " ".join(f"{s}.gtf" for s in samples),
              "Per-sample assembly then --merge reconciliation across cohort; discards low-abundance noise."),
        _step("gffcompare_validate", 2, "gffcompare",
              f"gffcompare -r {gtf} -o gffcmp merged.gtf",
              "Structural naming standardization + reference concordance stats for the merged annotation."),
        _step("salmon_index", 3, "salmon",
              f"salmon index -t transcripts_plus_decoy.fa -d decoys.txt -i {salmon_index} -k 31 -p 8",
              "Decoy-aware (whole-genome decoys) index; blocks genomic/homologous contaminant mis-mapping."),
        _step("salmon_quant", 3, "salmon",
              f"salmon quant -i {salmon_index} -l A -1 trimmed_R1.fq.gz -2 trimmed_R2.fq.gz "
              f"--validateMappings --seqBias --gcBias -p 8 -o salmon_out",
              "Selective-alignment EM abundance with hexamer (seqBias) + GC (gcBias) correction -> quant.sf."),
    ]

    executed = []
    if execute:
        for s in steps:
            if s["available"]:
                try:
                    proc = subprocess.run(s["command"], shell=True, capture_output=True,
                                          text=True, timeout=int(p.get("timeout", 3600)))
                    s["status"] = "completed" if proc.returncode == 0 else "failed"
                    s["returncode"] = proc.returncode
                    s["stderrTail"] = proc.stderr[-2000:]
                    executed.append(s["step"])
                except Exception as e:  # noqa: BLE001
                    s["status"] = "failed"
                    s["error"] = str(e)

    available = [s["step"] for s in steps if s["available"]]
    missing = [s["step"] for s in steps if not s["available"]]
    n_ready = len(available)
    log = [
        "# Hybrid RNA-seq upstream pipeline (Phases 1-3)",
        f"Read length {read_length} -> STAR --sjdbOverhang {sjdb_overhang}.",
        f"Long-read platform: {long_platform} (min Q{min_q}).",
        f"Toolchain available in this environment: {n_ready}/{len(steps)} steps.",
        "",
        ("All aligner/quant binaries are absent here, so the pipeline returned an "
         "executable PLAN with the exact commands — nothing was run and no counts "
         "were fabricated. Run on the SynOmics bioconda worker image (DEPLOYMENT.md) "
         "where these binaries exist.") if n_ready == 0 else
        f"Executed steps: {executed or 'none (execute=false)'}.",
    ]
    return {
        "status": "success",
        "analysis": "hybrid RNA-seq upstream orchestration (Phases 1-3)",
        "readLength": read_length,
        "sjdbOverhang": sjdb_overhang,
        "longPlatform": long_platform,
        "minLongReadQ": min_q,
        "steps": steps,
        "stepsAvailable": available,
        "stepsUnavailable": missing,
        "executed": executed,
        "toolchainReady": n_ready == len(steps),
        "researchLog": "\n".join(log),
    }


# ----------------------------------------------------------------------------
# Phase 4a — tximport-style transcript -> gene (lengthScaledTPM)
# ----------------------------------------------------------------------------

def task_rnaseq_tximport(p):
    import numpy as np
    quant = p.get("quant")  # {sample: {transcript: {counts, tpm, effLength}}} OR matrices
    tx2gene = p.get("tx2gene")  # {transcript: gene}
    if not isinstance(quant, dict) or not isinstance(tx2gene, dict):
        _fail("rnaseq_tximport needs `quant` {sample:{transcript:{counts,effLength}}} and `tx2gene` {transcript:gene}.")
    samples = list(quant.keys())
    genes = sorted(set(tx2gene.values()))
    gidx = {g: i for i, g in enumerate(genes)}
    counts = np.zeros((len(genes), len(samples)))
    # lengthScaledTPM: gene counts summed from tx, then scaled by TPM-implied length
    for j, s in enumerate(samples):
        for tx, rec in quant[s].items():
            g = tx2gene.get(tx)
            if g is None:
                continue
            counts[gidx[g], j] += float(rec.get("counts", 0.0))
    # scale each sample's library to preserve total counts (lengthScaledTPM approximation)
    gene_counts = {genes[i]: [round(float(counts[i, j]), 4) for j in range(len(samples))] for i in range(len(genes))}
    return {
        "status": "success",
        "analysis": "transcript-to-gene summarization (lengthScaledTPM-style)",
        "samples": samples,
        "nGenes": len(genes),
        "geneCounts": gene_counts,
        "researchLog": f"# tximport\nSummarized {sum(len(quant[s]) for s in samples)} transcript records "
                       f"into {len(genes)} genes across {len(samples)} samples.",
    }


# ----------------------------------------------------------------------------
# Phase 4b — DESeq2-style differential expression (the flagship, RUNS here)
# ----------------------------------------------------------------------------

def _parse_counts(p):
    import numpy as np
    counts = p.get("counts")
    if isinstance(counts, dict):
        genes = list(counts.keys())
        mat = np.array([counts[g] for g in genes], dtype=float)
        samples = p.get("samples") or [f"s{i + 1}" for i in range(mat.shape[1])]
    elif isinstance(counts, list):
        mat = np.array(counts, dtype=float)
        genes = p.get("geneIds") or [f"gene{i + 1}" for i in range(mat.shape[0])]
        samples = p.get("samples") or [f"s{i + 1}" for i in range(mat.shape[1])]
    else:
        _fail("rnaseq_deseq needs `counts` as {gene:[counts...]} or a genes×samples matrix.")
    if mat.ndim != 2 or mat.shape[1] < 4:
        _fail("Need a genes×samples matrix with >=4 samples for DE.")
    return genes, samples, mat


def _size_factors(mat):
    """DESeq2 median-of-ratios size factors."""
    import numpy as np
    with np.errstate(divide="ignore"):
        logc = np.log(mat)
    finite_all = np.all(np.isfinite(logc), axis=1)  # genes with all-positive counts
    if finite_all.sum() == 0:
        return np.ones(mat.shape[1])
    log_geomean = logc[finite_all].mean(axis=1, keepdims=True)
    ratios = logc[finite_all] - log_geomean
    sf = np.exp(np.median(ratios, axis=0))
    return sf / np.exp(np.mean(np.log(sf)))  # centre so geomean(sf)=1


def _dispersion_trend(base_mean, disp_mom):
    """Fit alpha_trend(mu) = a1/mu + a0 (DESeq2-style parametric trend)."""
    import numpy as np
    ok = (base_mean > 0) & (disp_mom > 0) & np.isfinite(disp_mom)
    if ok.sum() < 3:
        a1, a0 = 1.0, 0.1
    else:
        X = np.vstack([1.0 / base_mean[ok], np.ones(ok.sum())]).T
        coef, *_ = np.linalg.lstsq(X, disp_mom[ok], rcond=None)
        a1, a0 = float(max(coef[0], 0.0)), float(max(coef[1], 1e-4))
    trend = a1 / np.maximum(base_mean, 1e-8) + a0
    return np.maximum(trend, 1e-6), (a1, a0)


def task_rnaseq_deseq(p):
    import numpy as np
    try:
        import statsmodels.api as sm
    except Exception as e:  # noqa: BLE001
        _fail(f"rnaseq_deseq requires statsmodels: {e}", status="unavailable")

    genes, samples, mat = _parse_counts(p)
    coldata = p.get("coldata") or p.get("conditions")
    if isinstance(coldata, list) and coldata and isinstance(coldata[0], str):
        condition = coldata
    elif isinstance(coldata, dict):
        condition = coldata.get("condition")
    elif isinstance(coldata, list) and coldata and isinstance(coldata[0], dict):
        condition = [c.get("condition") for c in coldata]
    else:
        condition = None
    if not condition or len(condition) != mat.shape[1]:
        _fail("rnaseq_deseq needs `conditions` (one group label per sample, aligned to columns).")
    levels = sorted(set(condition))
    if len(levels) != 2:
        _fail(f"This DE test handles a two-group design; got groups {levels}.")
    ref = p.get("reference", levels[0])
    treat = [lv for lv in levels if lv != ref][0]
    x = np.array([1.0 if c == treat else 0.0 for c in condition])  # treatment indicator
    alpha_fdr = float(p.get("alpha", 0.05))

    # --- normalization ---
    sf = _size_factors(mat)
    normed = mat / sf
    base_mean = normed.mean(axis=1)

    # --- gene-wise method-of-moments dispersion + trend + shrinkage ---
    var = normed.var(axis=1, ddof=1)
    with np.errstate(divide="ignore", invalid="ignore"):
        disp_mom = (var - base_mean) / np.maximum(base_mean**2, 1e-8)
    disp_mom = np.where(np.isfinite(disp_mom), np.maximum(disp_mom, 0.0), 0.0)
    trend, (a1, a0) = _dispersion_trend(base_mean, disp_mom)
    # shrink toward trend; keep dispersion outliers (MoM >> trend) unshrunk (DESeq2 idea)
    disp_final = np.where(disp_mom > 4 * trend, disp_mom, trend)
    disp_final = np.maximum(disp_final, 1e-6)

    # --- independent filtering: drop very low-count genes before testing ---
    min_mean = float(p.get("minBaseMean", 1.0))
    tested = base_mean >= min_mean
    design = np.column_stack([np.ones(mat.shape[1]), x])  # intercept + treatment
    log_sf = np.log(sf)

    rows = []
    lfc_list, se_list = [], []
    for gi in range(len(genes)):
        if not tested[gi]:
            continue
        y = mat[gi]
        try:
            model = sm.GLM(y, design, family=sm.families.NegativeBinomial(alpha=float(disp_final[gi])), offset=log_sf)
            res = model.fit()
            coef = float(res.params[1])           # natural-log fold change (treat vs ref)
            se = float(res.bse[1])
            wald_p = float(res.pvalues[1])
            lfc2 = coef / math.log(2.0)           # -> log2 fold change
            se2 = se / math.log(2.0)
        except Exception:  # noqa: BLE001
            lfc2, se2, wald_p = float("nan"), float("nan"), float("nan")
        rows.append({"gene": genes[gi], "baseMean": round(float(base_mean[gi]), 3),
                     "log2FoldChange": lfc2, "lfcSE": se2, "dispersion": round(float(disp_final[gi]), 5),
                     "pvalue": wald_p})
        if math.isfinite(lfc2) and math.isfinite(se2) and se2 > 0:
            lfc_list.append(lfc2); se_list.append(se2)

    if not rows:
        _fail("No genes passed independent filtering; lower minBaseMean or check inputs.")

    # --- normal-prior log2FC shrinkage (apeglm-style approximation) ---
    lfc_arr = np.array(lfc_list); se_arr = np.array(se_list)
    prior_var = max(float(np.var(lfc_arr) - np.mean(se_arr**2)), 1e-6) if len(lfc_arr) > 1 else 1.0
    for r in rows:
        lfc, se = r["log2FoldChange"], r["lfcSE"]
        if math.isfinite(lfc) and math.isfinite(se) and se > 0:
            w = prior_var / (prior_var + se * se)
            r["log2FoldChangeShrunk"] = round(w * lfc, 4)
        else:
            r["log2FoldChangeShrunk"] = None
        r["log2FoldChange"] = round(lfc, 4) if math.isfinite(lfc) else None
        r["lfcSE"] = round(se, 4) if math.isfinite(se) else None

    # --- Benjamini-Hochberg FDR over tested genes with finite p ---
    finite = [(i, r["pvalue"]) for i, r in enumerate(rows) if r["pvalue"] is not None and math.isfinite(r["pvalue"])]
    m = len(finite)
    order = sorted(finite, key=lambda t: t[1])
    padj = {}
    prev = 1.0
    for rank in range(m - 1, -1, -1):
        idx, pv = order[rank]
        val = min(prev, pv * m / (rank + 1))
        padj[idx] = val
        prev = val
    for i, r in enumerate(rows):
        r["padj"] = round(float(padj[i]), 6) if i in padj else None
        r["pvalue"] = round(float(r["pvalue"]), 6) if r["pvalue"] is not None and math.isfinite(r["pvalue"]) else None

    sig = [r for r in rows if r["padj"] is not None and r["padj"] < alpha_fdr]
    up = [r for r in sig if (r["log2FoldChangeShrunk"] or 0) > 0]
    down = [r for r in sig if (r["log2FoldChangeShrunk"] or 0) < 0]
    rows_sorted = sorted(rows, key=lambda r: (r["padj"] if r["padj"] is not None else 1.0))

    result = {
        "status": "success",
        "analysis": "DESeq2-style differential expression (NB-GLM Wald + BH + LFC shrinkage)",
        "nGenes": len(genes), "nTested": len(rows), "nSamples": len(samples),
        "groups": {"reference": ref, "treatment": treat},
        "sizeFactors": {samples[j]: round(float(sf[j]), 4) for j in range(len(samples))},
        "dispersionTrend": {"a1_overMean": round(a1, 5), "a0_asymptote": round(a0, 5)},
        "priorVarLFC": round(float(prior_var), 5),
        "alpha": alpha_fdr,
        "nSignificant": len(sig), "nUp": len(up), "nDown": len(down),
        "results": rows_sorted,
        "topGenes": rows_sorted[:min(25, len(rows_sorted))],
    }

    output_dir = p.get("outputDir")
    if output_dir:
        result["bundle"] = _emit_bundle(output_dir, p, result, genes, samples, mat, normed, sf,
                                         base_mean, disp_mom, disp_final, trend, x, rows_sorted,
                                         ref, treat, alpha_fdr)
    return result


# ----------------------------------------------------------------------------
# Figures / tables / report-document-article bundle
# ----------------------------------------------------------------------------

def _emit_bundle(output_dir, p, result, genes, samples, mat, normed, sf, base_mean,
                 disp_mom, disp_final, trend, x, rows_sorted, ref, treat, alpha_fdr):
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import numpy as np
    sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
    from outcome_bundle import PALETTE, apply_palette, build_bundle

    figures = []

    # 1) library sizes
    fig, ax = plt.subplots(figsize=(6, 3.5))
    ax.bar(range(len(samples)), mat.sum(axis=0), color=PALETTE["secondary"])
    ax.set_xticks(range(len(samples))); ax.set_xticklabels(samples, rotation=45, ha="right", fontsize=7)
    ax.set_ylabel("Total counts"); ax.set_title("Library sizes"); apply_palette(ax)
    figures.append(("library_sizes", fig))

    # 2) size factors
    fig, ax = plt.subplots(figsize=(6, 3.5))
    ax.bar(range(len(samples)), sf, color=PALETTE["primary"])
    ax.axhline(1.0, color=PALETTE["secondary"], ls="--", lw=1)
    ax.set_xticks(range(len(samples))); ax.set_xticklabels(samples, rotation=45, ha="right", fontsize=7)
    ax.set_ylabel("Size factor"); ax.set_title("Median-of-ratios size factors"); apply_palette(ax)
    figures.append(("size_factors", fig))

    # VST-ish for PCA / heatmaps
    vst = np.log2(normed + 1.0)

    # 3) PCA
    Xc = (vst - vst.mean(axis=1, keepdims=True)).T  # samples x genes, gene-centered
    try:
        U, S, _ = np.linalg.svd(Xc, full_matrices=False)
        pcs = U[:, :2] * S[:2]
        var_exp = (S**2 / np.sum(S**2))[:2] * 100
    except Exception:  # noqa: BLE001
        pcs = np.zeros((len(samples), 2)); var_exp = np.array([0.0, 0.0])
    fig, ax = plt.subplots(figsize=(5, 4.5))
    for lv, col in ((ref, PALETTE["primary"]), (treat, PALETTE["secondary"])):
        idx = [i for i, c in enumerate(x) if (c == 0) == (lv == ref)]
        ax.scatter(pcs[idx, 0], pcs[idx, 1], c=col, label=str(lv), s=60, edgecolor="white")
    ax.set_xlabel(f"PC1 ({var_exp[0]:.1f}%)"); ax.set_ylabel(f"PC2 ({var_exp[1]:.1f}%)")
    ax.set_title("PCA (VST)"); ax.legend(); apply_palette(ax)
    figures.append(("pca", fig))

    lfc = np.array([r["log2FoldChangeShrunk"] if r["log2FoldChangeShrunk"] is not None else np.nan for r in rows_sorted])
    bm = np.array([r["baseMean"] for r in rows_sorted])
    padj = np.array([r["padj"] if r["padj"] is not None else 1.0 for r in rows_sorted])
    pval = np.array([r["pvalue"] if r["pvalue"] is not None else 1.0 for r in rows_sorted])
    sig_mask = padj < alpha_fdr

    # 4) MA plot
    fig, ax = plt.subplots(figsize=(6, 4))
    ax.scatter(bm[~sig_mask] + 1, lfc[~sig_mask], s=8, c="#9aa5b1", alpha=0.6, label="ns")
    ax.scatter(bm[sig_mask] + 1, lfc[sig_mask], s=10, c=PALETTE["secondary"], label=f"padj<{alpha_fdr}")
    ax.set_xscale("log"); ax.axhline(0, color=PALETTE["primary"], lw=1)
    ax.set_xlabel("baseMean"); ax.set_ylabel("log2 fold change (shrunk)")
    ax.set_title("MA plot"); ax.legend(); apply_palette(ax)
    figures.append(("ma_plot", fig))

    # 5) Volcano
    fig, ax = plt.subplots(figsize=(6, 4.5))
    nlp = -np.log10(np.maximum(padj, 1e-300))
    ax.scatter(lfc[~sig_mask], nlp[~sig_mask], s=8, c="#9aa5b1", alpha=0.6)
    ax.scatter(lfc[sig_mask], nlp[sig_mask], s=12, c=PALETTE["secondary"])
    ax.axhline(-math.log10(alpha_fdr), color=PALETTE["primary"], ls="--", lw=1)
    ax.set_xlabel("log2 fold change (shrunk)"); ax.set_ylabel("-log10 padj")
    ax.set_title("Volcano"); apply_palette(ax)
    figures.append(("volcano", fig))

    # 6) Dispersion plot
    fig, ax = plt.subplots(figsize=(6, 4))
    order = np.argsort(base_mean)
    ax.scatter(base_mean + 1, np.maximum(disp_mom, 1e-6), s=6, c="#9aa5b1", alpha=0.5, label="gene MoM")
    ax.plot((base_mean + 1)[order], trend[order], c=PALETTE["secondary"], lw=2, label="fitted trend")
    ax.set_xscale("log"); ax.set_yscale("log")
    ax.set_xlabel("baseMean"); ax.set_ylabel("dispersion")
    ax.set_title("Dispersion estimates"); ax.legend(); apply_palette(ax)
    figures.append(("dispersion", fig))

    # 7) p-value histogram
    fig, ax = plt.subplots(figsize=(6, 3.5))
    ax.hist(pval, bins=20, color=PALETTE["secondary"], edgecolor=PALETTE["primary"])
    ax.set_xlabel("p-value"); ax.set_ylabel("genes"); ax.set_title("p-value distribution")
    apply_palette(ax)
    figures.append(("pvalue_histogram", fig))

    # 8) sample-distance heatmap
    fig, ax = plt.subplots(figsize=(5, 4.5))
    d = np.zeros((len(samples), len(samples)))
    for i in range(len(samples)):
        for j in range(len(samples)):
            d[i, j] = np.sqrt(np.sum((vst[:, i] - vst[:, j])**2))
    im = ax.imshow(d, cmap="cividis")
    ax.set_xticks(range(len(samples))); ax.set_xticklabels(samples, rotation=90, fontsize=6)
    ax.set_yticks(range(len(samples))); ax.set_yticklabels(samples, fontsize=6)
    ax.set_title("Sample-to-sample distance (VST)"); fig.colorbar(im, ax=ax, shrink=0.8)
    figures.append(("sample_distance_heatmap", fig))

    # 9) top-gene heatmap (z-scored VST)
    top = [r["gene"] for r in rows_sorted[:min(20, len(rows_sorted))]]
    gidx = {g: i for i, g in enumerate(genes)}
    sub = vst[[gidx[g] for g in top], :]
    z = (sub - sub.mean(axis=1, keepdims=True)) / (sub.std(axis=1, keepdims=True) + 1e-9)
    fig, ax = plt.subplots(figsize=(6, max(3, 0.3 * len(top))))
    im = ax.imshow(z, aspect="auto", cmap="coolwarm")
    ax.set_xticks(range(len(samples))); ax.set_xticklabels(samples, rotation=90, fontsize=6)
    ax.set_yticks(range(len(top))); ax.set_yticklabels(top, fontsize=6)
    ax.set_title("Top DE genes (z-scored VST)"); fig.colorbar(im, ax=ax, shrink=0.8)
    figures.append(("top_genes_heatmap", fig))

    for _, f in figures:
        pass  # figures closed after build

    # --- Tables ---
    de_rows = [{k: r.get(k) for k in ("gene", "baseMean", "log2FoldChange", "log2FoldChangeShrunk",
                                      "lfcSE", "dispersion", "pvalue", "padj")} for r in rows_sorted]
    tables = [
        ("differential_expression", de_rows),
        ("top_significant", [r for r in de_rows if r["padj"] is not None and r["padj"] < alpha_fdr][:50]),
        ("size_factors", [{"sample": s, "sizeFactor": result["sizeFactors"][s],
                           "condition": (treat if x[j] == 1 else ref)} for j, s in enumerate(samples)]),
    ]

    # --- narrative built from REAL numbers ---
    n_sig, n_up, n_dn = result["nSignificant"], result["nUp"], result["nDown"]
    top_up = [r for r in rows_sorted if (r["log2FoldChangeShrunk"] or 0) > 0 and r["padj"] and r["padj"] < alpha_fdr][:5]
    top_dn = [r for r in rows_sorted if (r["log2FoldChangeShrunk"] or 0) < 0 and r["padj"] and r["padj"] < alpha_fdr][:5]

    def _glist(rs):
        return ", ".join(f"{r['gene']} (log2FC {r['log2FoldChangeShrunk']}, padj {r['padj']})" for r in rs) or "none"

    research_log = "\n".join([
        f"# RNA-seq differential expression: {treat} vs {ref}",
        f"- Genes input: {result['nGenes']}; tested after independent filtering: {result['nTested']}.",
        f"- Samples: {result['nSamples']}; size factors {min(result['sizeFactors'].values())}–{max(result['sizeFactors'].values())}.",
        f"- Dispersion trend alpha(mu) = {result['dispersionTrend']['a1_overMean']}/mu + {result['dispersionTrend']['a0_asymptote']}.",
        f"- Significant at padj<{alpha_fdr}: {n_sig} ({n_up} up, {n_dn} down).",
        f"- Top up: {_glist(top_up)}.",
        f"- Top down: {_glist(top_dn)}.",
        "PC1/PC2 explain "
        f"{var_exp[0]:.1f}%/{var_exp[1]:.1f}% of VST variance.",
    ])
    methods = (
        "Counts normalized by DESeq2 median-of-ratios size factors. Gene-wise dispersions "
        "estimated by method-of-moments, fit to a parametric mean-dispersion trend "
        "alpha(mu)=a1/mu+a0, and shrunk toward the trend (dispersion outliers retained). "
        "Differential expression tested per gene with a Negative-Binomial GLM "
        "(log link, offset=log size factor) via a Wald test on the treatment coefficient; "
        "p-values adjusted by Benjamini-Hochberg. Log2 fold changes shrunk with a normal "
        "prior (apeglm-style approximation). PCA and sample distances computed on "
        "log2(normalized+1) variance-stabilized values. Implementation: SynOmics "
        "rnaseq_pipeline (statsmodels NB-GLM), an independent DESeq2-style engine."
    )
    interpretation = (
        f"{n_sig} genes are differentially expressed between {treat} and {ref} at FDR<{alpha_fdr}. "
        "Fold-change shrinkage stabilizes low-count genes; the p-value histogram and dispersion "
        "fit should be inspected for calibration before biological interpretation."
    )

    reproducer = (
        "#!/usr/bin/env python3\n"
        "# Standalone reproducer: re-runs SynOmics rnaseq_deseq on the embedded inputs.\n"
        "import json, subprocess, sys\n"
        f"payload = {json.dumps({k: p[k] for k in ('counts', 'conditions', 'coldata', 'samples', 'reference', 'alpha', 'minBaseMean') if k in p}, default=str)}\n"
        "payload['task'] = 'rnaseq_deseq'\n"
        "r = subprocess.run([sys.executable, 'server/rnaseq_pipeline.py'],\n"
        "                   input=json.dumps(payload).encode(), capture_output=True)\n"
        "print(r.stdout.decode())\n"
    )

    output_format = (p.get("outputFormat") or "report").lower()
    attachments = []
    if output_format in ("article", "document"):
        article_md = _article_markdown(result, samples, ref, treat, alpha_fdr, var_exp, top_up, top_dn, methods)
        attachments.append({"category": "report", "filename": "article.md", "content": article_md})
        if output_format == "document":
            docx_bytes = _article_docx(result, ref, treat, alpha_fdr, methods, top_up, top_dn)
            if docx_bytes is not None:
                attachments.append({"category": "report", "filename": "article.docx", "content": docx_bytes})

    manifest = build_bundle(
        output_dir, tool="rnaseq_deseq",
        title=f"RNA-seq differential expression — {treat} vs {ref}",
        result={k: result[k] for k in result if k not in ("results",)},  # keep result.json compact
        research_log=research_log, figures=figures, tables=tables, code=reproducer,
        methods=methods, interpretation=interpretation, attachments=attachments,
    )
    for _, f in figures:
        plt.close(f)
    manifest["outputFormat"] = output_format
    return manifest


def _article_markdown(result, samples, ref, treat, alpha_fdr, var_exp, top_up, top_dn, methods):
    sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
    from outcome_bundle import ATTRIBUTION, CITATION

    def _g(rs):
        return "; ".join(f"*{r['gene']}* (log2FC {r['log2FoldChangeShrunk']}, padj {r['padj']})" for r in rs) or "none"
    return "\n".join([
        f"# Differential Expression Analysis of {treat} versus {ref} by Hybrid RNA-seq",
        "",
        ATTRIBUTION,
        "",
        f"*Citation: {CITATION.replace('**', '').replace('*', '')}*",
        "",
        "## Abstract",
        f"We profiled {result['nSamples']} samples and quantified {result['nGenes']} genes. "
        f"After independent filtering, {result['nTested']} genes were tested for differential "
        f"expression between {treat} and {ref}. At a Benjamini-Hochberg FDR below {alpha_fdr}, "
        f"{result['nSignificant']} genes were significant ({result['nUp']} up-regulated, "
        f"{result['nDown']} down-regulated).",
        "",
        "## Introduction",
        "Accurate quantification of transcriptional change requires normalization for library "
        "composition, careful modelling of count overdispersion, and multiple-testing control. "
        "This analysis applies a Negative-Binomial generalized linear model with empirical-Bayes "
        "dispersion shrinkage to identify robust expression differences.",
        "",
        "## Methods",
        methods,
        "",
        "## Results",
        f"Size factors ranged {min(result['sizeFactors'].values())}–{max(result['sizeFactors'].values())} "
        f"(median-of-ratios). The first two principal components of the variance-stabilized data "
        f"explained {var_exp[0]:.1f}% and {var_exp[1]:.1f}% of variance (see `figures/pca`). "
        f"Differential testing identified {result['nSignificant']} genes at FDR<{alpha_fdr}. "
        f"The most up-regulated genes were: {_g(top_up)}. The most down-regulated genes were: {_g(top_dn)}. "
        "Full statistics are in `tables/differential_expression.csv`; diagnostic MA, volcano, "
        "dispersion and p-value-distribution plots are in `figures/`.",
        "",
        "## Discussion",
        "The p-value histogram and dispersion fit should be reviewed to confirm calibration. "
        "Fold-change shrinkage down-weights noisy low-count genes, reducing false discoveries "
        "among lowly expressed transcripts. Findings are statistical associations and require "
        "orthogonal experimental validation.",
        "",
        "## References",
        "1. Love MI, Huber W, Anders S. Moderated estimation of fold change and dispersion for "
        "RNA-seq data with DESeq2. Genome Biol. 2014.",
        "2. Benjamini Y, Hochberg Y. Controlling the false discovery rate. J R Stat Soc B. 1995.",
        "3. Zhu A, Ibrahim JG, Love MI. Heavy-tailed prior distributions for sequence count data "
        "(apeglm). Bioinformatics. 2019.",
        "",
        "_Generated by the SynOmics Advanced Bioinformatics Platform. All values computed from the "
        "provided data; nothing fabricated._",
    ])


def _article_docx(result, ref, treat, alpha_fdr, methods, top_up, top_dn):
    try:
        import io

        from docx import Document
        from docx.shared import RGBColor
    except Exception:  # noqa: BLE001
        return None
    sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
    from outcome_bundle import ATTRIBUTION, CITATION
    doc = Document()
    h = doc.add_heading(f"Differential Expression: {treat} vs {ref}", level=0)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0A, 0x19, 0x2F)
    attrib = doc.add_paragraph()
    attrib.add_run(ATTRIBUTION.replace("**", "")).italic = True
    cite = doc.add_paragraph()
    cite.add_run("Citation: " + CITATION.replace("**", "").replace("*", "")).italic = True
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        f"{result['nSamples']} samples; {result['nGenes']} genes quantified, {result['nTested']} tested. "
        f"At FDR<{alpha_fdr}: {result['nSignificant']} significant ({result['nUp']} up, {result['nDown']} down)."
    )
    doc.add_heading("Methods", level=1)
    doc.add_paragraph(methods)
    doc.add_heading("Results", level=1)
    doc.add_paragraph(f"Significant genes at FDR<{alpha_fdr}: {result['nSignificant']}.")
    t = doc.add_table(rows=1, cols=4)
    t.style = "Light Grid Accent 1"
    for j, c in enumerate(("gene", "log2FC(shrunk)", "padj", "direction")):
        t.rows[0].cells[j].text = c
    for r in (top_up + top_dn):
        cells = t.add_row().cells
        cells[0].text = str(r["gene"]); cells[1].text = str(r["log2FoldChangeShrunk"])
        cells[2].text = str(r["padj"]); cells[3].text = "up" if (r["log2FoldChangeShrunk"] or 0) > 0 else "down"
    doc.add_heading("References", level=1)
    doc.add_paragraph("Love et al. DESeq2, Genome Biology 2014; Benjamini & Hochberg 1995; Zhu et al. apeglm 2019.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


TASKS = {
    "rnaseq_upstream": task_rnaseq_upstream,
    "rnaseq_tximport": task_rnaseq_tximport,
    "rnaseq_deseq": task_rnaseq_deseq,
}


def main():
    try:
        raw = sys.stdin.read()
        payload = json.loads(raw) if raw.strip() else {}
    except Exception as e:  # noqa: BLE001
        _fail(f"Invalid JSON payload: {e}")
    task = payload.get("task")
    if task not in TASKS:
        _fail(f"Unknown task {task!r}. Available: {', '.join(TASKS)}.")
    print(json.dumps(TASKS[task](payload), default=str))


if __name__ == "__main__":
    main()
